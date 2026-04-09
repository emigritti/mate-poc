"""
Unit tests — Requirements parsing (routers/requirements.py).
CLAUDE.md §7: Business logic, edge cases, validation.

Coverage:
  _parse_markdown() — structured (bullet) mode:
  - Valid frontmatter → correct source/target
  - Mandatory section heading → mandatory=True
  - Non-mandatory / Optional heading → mandatory=False
  - Mixed sections → correct mandatory flags per item
  - Bullet formats: 3-part (REQ-ID|Cat|Desc), 2-part (Cat|Desc), 1-part (Desc)
  - No frontmatter → fallback filename parsing ("erp-to-salsify.md")
  - No frontmatter, unparseable filename → source/target = "Unknown"
  - Empty body → empty list
  - Skips non-bullet lines in sections (structured mode)

  _parse_markdown() — unstructured (prose) mode:
  - H1 + H2 → each H2 section = one requirement
  - Only H1 → each H1 section = one requirement
  - No headings → each paragraph = one requirement
  - Parent H1 "mandatory" heading propagates mandatory flag to H2 children
  - H2 heading text used as category; prose body as description

  _parse_prose_requirements() / _parse_paragraphs_as_requirements():
  - Direct unit tests for prose grouping logic

  _parse_docx_requirements():
  - Heading 1 + Heading 2 → each H2 section = one requirement
  - Only Heading 1 → each H1 section = one requirement
  - No heading styles → each paragraph = one requirement
  - Returns "Unknown" for source/target (user fills via validation modal)

  _parse_csv():
  - Mandatory column "true" / "yes" / "1" → mandatory=True
  - Absent Mandatory column → mandatory=False (default)

  Upload endpoint:
  - .md file accepted (200)
  - .md with text/plain MIME accepted (200)
  - .txt file accepted (200)
  - Parsed mandatory flags visible via GET /requirements
"""

import io
import pytest
from unittest.mock import AsyncMock, patch, MagicMock

from fastapi.testclient import TestClient

# ── Parser unit tests (no HTTP) ───────────────────────────────────────────────

from routers.requirements import (
    _parse_markdown,
    _parse_csv,
    _parse_prose_requirements,
    _parse_paragraphs_as_requirements,
    _parse_docx_requirements,
    _is_unstructured,
)


_SAMPLE_MD = """\
---
source: ERP
target: Salsify
---

## Mandatory Requirements

- REQ-M01 | Product Collection | Sync daily articles from ERP

## Non-Mandatory Requirements

- REQ-O01 | Reporting | Generate weekly report
"""

_SAMPLE_MD_OPTIONAL_HEADING = """\
---
source: PLM
target: DAM
---

## Mandatory

- REQ-M01 | Sync | Transfer product data

## Optional

- REQ-O01 | Archive | Archive old assets
"""


class TestParseMarkdownFrontmatter:
    def test_extracts_source_from_frontmatter(self):
        reqs = _parse_markdown(_SAMPLE_MD)
        assert all(r.source_system == "ERP" for r in reqs)

    def test_extracts_target_from_frontmatter(self):
        reqs = _parse_markdown(_SAMPLE_MD)
        assert all(r.target_system == "Salsify" for r in reqs)

    def test_fallback_filename_dash_to(self):
        md = "## Mandatory\n- REQ-1 | Cat | Desc\n"
        reqs = _parse_markdown(md, filename="erp-to-salsify.md")
        assert reqs[0].source_system == "erp"
        assert reqs[0].target_system == "salsify"

    def test_fallback_filename_arrow(self):
        md = "## Mandatory\n- REQ-1 | Cat | Desc\n"
        reqs = _parse_markdown(md, filename="plm→dam.md")
        assert reqs[0].source_system == "plm"
        assert reqs[0].target_system == "dam"

    def test_fallback_unknown_when_no_separator(self):
        md = "## Mandatory\n- REQ-1 | Cat | Desc\n"
        reqs = _parse_markdown(md, filename="requirements.md")
        assert reqs[0].source_system == "Unknown"
        assert reqs[0].target_system == "Unknown"


class TestParseMarkdownMandatory:
    def test_mandatory_section_sets_flag_true(self):
        reqs = _parse_markdown(_SAMPLE_MD)
        mandatory_reqs = [r for r in reqs if r.req_id == "REQ-M01"]
        assert len(mandatory_reqs) == 1
        assert mandatory_reqs[0].mandatory is True

    def test_non_mandatory_section_sets_flag_false(self):
        reqs = _parse_markdown(_SAMPLE_MD)
        optional_reqs = [r for r in reqs if r.req_id == "REQ-O01"]
        assert len(optional_reqs) == 1
        assert optional_reqs[0].mandatory is False

    def test_optional_heading_sets_mandatory_false(self):
        reqs = _parse_markdown(_SAMPLE_MD_OPTIONAL_HEADING)
        optional_reqs = [r for r in reqs if r.req_id == "REQ-O01"]
        assert optional_reqs[0].mandatory is False

    def test_mandatory_heading_without_non_prefix_sets_true(self):
        reqs = _parse_markdown(_SAMPLE_MD_OPTIONAL_HEADING)
        mandatory_reqs = [r for r in reqs if r.req_id == "REQ-M01"]
        assert mandatory_reqs[0].mandatory is True

    def test_mixed_sections_correct_count(self):
        reqs = _parse_markdown(_SAMPLE_MD)
        assert sum(1 for r in reqs if r.mandatory) == 1
        assert sum(1 for r in reqs if not r.mandatory) == 1


class TestParseMarkdownBulletFormats:
    def test_three_part_bullet_uses_explicit_req_id(self):
        md = "---\nsource: A\ntarget: B\n---\n## Mandatory\n- REQ-X01 | Category | Some description\n"
        reqs = _parse_markdown(md)
        assert reqs[0].req_id == "REQ-X01"
        assert reqs[0].category == "Category"
        assert reqs[0].description == "Some description"

    def test_two_part_bullet_auto_generates_req_id(self):
        md = "---\nsource: A\ntarget: B\n---\n## Mandatory\n- Sync | Transfer product data\n"
        reqs = _parse_markdown(md)
        assert reqs[0].req_id.startswith("R-")
        assert reqs[0].category == "Sync"
        assert reqs[0].description == "Transfer product data"

    def test_one_part_bullet_sets_general_category(self):
        md = "---\nsource: A\ntarget: B\n---\n## Mandatory\n- Just a plain description\n"
        reqs = _parse_markdown(md)
        assert reqs[0].req_id.startswith("R-")
        assert reqs[0].category == "General"
        assert reqs[0].description == "Just a plain description"

    def test_non_bullet_lines_ignored(self):
        md = "---\nsource: A\ntarget: B\n---\n## Mandatory\nSome prose text.\nMore prose.\n- REQ-1 | Cat | Desc\n"
        reqs = _parse_markdown(md)
        assert len(reqs) == 1

    def test_empty_markdown_returns_empty_list(self):
        reqs = _parse_markdown("---\nsource: A\ntarget: B\n---\n")
        assert reqs == []


class TestParseCsvMandatoryColumn:
    def test_mandatory_true_string(self):
        csv_text = "ReqID,Source,Target,Category,Description,Mandatory\nR-1,ERP,Salsify,Sync,Desc,true\n"
        reqs = _parse_csv(csv_text)
        assert reqs[0].mandatory is True

    def test_mandatory_yes_string(self):
        csv_text = "ReqID,Source,Target,Category,Description,Mandatory\nR-1,ERP,Salsify,Sync,Desc,yes\n"
        reqs = _parse_csv(csv_text)
        assert reqs[0].mandatory is True

    def test_mandatory_1_string(self):
        csv_text = "ReqID,Source,Target,Category,Description,Mandatory\nR-1,ERP,Salsify,Sync,Desc,1\n"
        reqs = _parse_csv(csv_text)
        assert reqs[0].mandatory is True

    def test_mandatory_false_when_column_absent(self):
        csv_text = "ReqID,Source,Target,Category,Description\nR-1,ERP,Salsify,Sync,Desc\n"
        reqs = _parse_csv(csv_text)
        assert reqs[0].mandatory is False

    def test_mandatory_false_for_empty_value(self):
        csv_text = "ReqID,Source,Target,Category,Description,Mandatory\nR-1,ERP,Salsify,Sync,Desc,\n"
        reqs = _parse_csv(csv_text)
        assert reqs[0].mandatory is False


# ── Upload endpoint tests ─────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def client():
    with (
        patch("db.init_db",          new_callable=AsyncMock),
        patch("db.close_db",         new_callable=AsyncMock),
        patch("main._init_chromadb", new_callable=AsyncMock),
    ):
        from main import app
        with TestClient(app) as c:
            yield c


_SAMPLE_MD_BYTES = _SAMPLE_MD.encode("utf-8")


class TestUploadMarkdown:
    def test_md_file_returns_200(self, client):
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.md", io.BytesIO(_SAMPLE_MD_BYTES), "text/markdown")},
        )
        assert res.status_code == 200

    def test_md_file_parsed_count(self, client):
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.md", io.BytesIO(_SAMPLE_MD_BYTES), "text/markdown")},
        )
        assert res.json()["total_parsed"] == 2

    def test_md_file_with_plain_mime_accepted(self, client):
        # Browsers may send text/plain for .md — filename extension takes precedence
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.md", io.BytesIO(_SAMPLE_MD_BYTES), "text/plain")},
        )
        assert res.status_code == 200

    def test_md_preview_contains_one_pair(self, client):
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.md", io.BytesIO(_SAMPLE_MD_BYTES), "text/markdown")},
        )
        preview = res.json()["preview"]
        assert len(preview) == 1
        assert preview[0]["source"] == "ERP"
        assert preview[0]["target"] == "Salsify"

    def test_mandatory_flags_visible_in_get(self, client):
        client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.md", io.BytesIO(_SAMPLE_MD_BYTES), "text/markdown")},
        )
        reqs = client.get("/api/v1/requirements").json()["data"]
        mandatory_flags = {r["req_id"]: r["mandatory"] for r in reqs}
        assert mandatory_flags.get("REQ-M01") is True
        assert mandatory_flags.get("REQ-O01") is False

    def test_unsupported_file_type_still_rejected(self, client):
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("data.bin", io.BytesIO(b"\x00\x01"), "application/octet-stream")},
        )
        assert res.status_code == 415

    def test_txt_file_accepted(self, client):
        txt = b"## Feature A\nSync products daily.\n\n## Feature B\nArchive old assets.\n"
        res = client.post(
            "/api/v1/requirements/upload",
            files={"file": ("reqs.txt", io.BytesIO(txt), "text/plain")},
        )
        assert res.status_code == 200
        assert res.json()["total_parsed"] == 2


# ── _is_unstructured ──────────────────────────────────────────────────────────

class TestIsUnstructured:
    def test_bullet_body_is_structured(self):
        assert _is_unstructured("## Section\n- item one\n- item two\n") is False

    def test_pure_prose_is_unstructured(self):
        assert _is_unstructured("## Section\nSome prose text.\nMore prose.\n") is True

    def test_mixed_prefers_structured(self):
        # Even one bullet → structured
        assert _is_unstructured("Prose line.\n- bullet\n") is False

    def test_empty_body_is_unstructured(self):
        assert _is_unstructured("") is True


# ── _parse_prose_requirements ─────────────────────────────────────────────────

_PROSE_H1_H2 = """\
# Mandatory Requirements

## Product Sync
Sync all products from ERP to Salsify daily.
Only active SKUs should be included.

## Asset Archive
Archive obsolete digital assets to cold storage.

# Non-Mandatory Requirements

## Reporting
Generate a weekly sync status report.
"""

_PROSE_H1_ONLY = """\
## Product Sync
Transfer product master data from ERP.

## Asset Management
Manage digital assets lifecycle.
"""

_PROSE_NO_HEADINGS = """\
Transfer product master data from ERP.

Archive obsolete digital assets.

Generate a weekly sync status report.
"""


class TestParseProseRequirements:
    def test_h1_h2_creates_one_req_per_h2(self):
        reqs = _parse_prose_requirements(_PROSE_H1_H2, "ERP", "Salsify")
        assert len(reqs) == 3

    def test_h2_heading_becomes_category(self):
        reqs = _parse_prose_requirements(_PROSE_H1_H2, "ERP", "Salsify")
        categories = {r.category for r in reqs}
        assert "Product Sync" in categories
        assert "Asset Archive" in categories
        assert "Reporting" in categories

    def test_prose_body_becomes_description(self):
        reqs = _parse_prose_requirements(_PROSE_H1_H2, "ERP", "Salsify")
        sync_req = next(r for r in reqs if r.category == "Product Sync")
        assert "Sync all products" in sync_req.description
        assert "Only active SKUs" in sync_req.description

    def test_mandatory_flag_from_parent_h1(self):
        reqs = _parse_prose_requirements(_PROSE_H1_H2, "ERP", "Salsify")
        sync_req = next(r for r in reqs if r.category == "Product Sync")
        report_req = next(r for r in reqs if r.category == "Reporting")
        assert sync_req.mandatory is True
        assert report_req.mandatory is False

    def test_h1_only_creates_one_req_per_h1(self):
        reqs = _parse_prose_requirements(_PROSE_H1_ONLY, "A", "B")
        assert len(reqs) == 2

    def test_h1_only_heading_becomes_category(self):
        reqs = _parse_prose_requirements(_PROSE_H1_ONLY, "A", "B")
        assert reqs[0].category == "Product Sync"
        assert reqs[1].category == "Asset Management"

    def test_source_target_propagated(self):
        reqs = _parse_prose_requirements(_PROSE_H1_ONLY, "ERP", "DAM")
        assert all(r.source_system == "ERP" for r in reqs)
        assert all(r.target_system == "DAM" for r in reqs)

    def test_auto_req_id_generated(self):
        reqs = _parse_prose_requirements(_PROSE_H1_ONLY, "A", "B")
        assert all(r.req_id.startswith("R-") for r in reqs)

    def test_empty_body_returns_empty_list(self):
        assert _parse_prose_requirements("", "A", "B") == []


class TestParseParagraphsAsRequirements:
    def test_blank_lines_separate_requirements(self):
        reqs = _parse_paragraphs_as_requirements(_PROSE_NO_HEADINGS, "A", "B")
        assert len(reqs) == 3

    def test_paragraph_text_becomes_description(self):
        reqs = _parse_paragraphs_as_requirements(_PROSE_NO_HEADINGS, "A", "B")
        assert "Transfer product master data" in reqs[0].description

    def test_category_is_general(self):
        reqs = _parse_paragraphs_as_requirements(_PROSE_NO_HEADINGS, "A", "B")
        assert all(r.category == "General" for r in reqs)

    def test_mandatory_defaults_false(self):
        reqs = _parse_paragraphs_as_requirements(_PROSE_NO_HEADINGS, "A", "B")
        assert all(r.mandatory is False for r in reqs)


# ── _parse_markdown unstructured mode ────────────────────────────────────────

class TestParseMarkdownUnstructured:
    def test_prose_md_routes_to_prose_parser(self):
        md = "---\nsource: ERP\ntarget: Salsify\n---\n\n## Product Sync\nSync products daily.\n"
        reqs = _parse_markdown(md)
        assert len(reqs) == 1
        assert reqs[0].category == "Product Sync"
        assert "Sync products daily" in reqs[0].description

    def test_frontmatter_source_target_respected_in_prose_mode(self):
        md = "---\nsource: PLM\ntarget: DAM\n---\n\n## Feature\nSome capability.\n"
        reqs = _parse_markdown(md)
        assert reqs[0].source_system == "PLM"
        assert reqs[0].target_system == "DAM"

    def test_prose_with_mandatory_parent(self):
        md = (
            "---\nsource: A\ntarget: B\n---\n\n"
            "# Mandatory Requirements\n\n"
            "## Sync\nSync data.\n\n"
            "# Optional\n\n"
            "## Reporting\nGenerate report.\n"
        )
        reqs = _parse_markdown(md)
        assert len(reqs) == 2
        sync = next(r for r in reqs if r.category == "Sync")
        report = next(r for r in reqs if r.category == "Reporting")
        assert sync.mandatory is True
        assert report.mandatory is False

    def test_txt_filename_fallback_parsing(self):
        txt = "## Feature A\nSync products.\n"
        reqs = _parse_markdown(txt, filename="erp-to-salsify.txt")
        assert reqs[0].source_system == "erp"
        assert reqs[0].target_system == "salsify"


# ── _parse_docx_requirements ──────────────────────────────────────────────────

def _make_docx_bytes(paragraphs: list[tuple[str | None, str]]) -> bytes:
    """Build a minimal .docx in memory.

    paragraphs: list of (style_name, text) where style_name is
    'Heading 1', 'Heading 2', or None (Normal paragraph).
    """
    from docx import Document as DocxDocument
    doc = DocxDocument()
    for style, text in paragraphs:
        if style:
            doc.add_heading(text, level=int(style.split()[-1]))
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class TestParseDocxRequirements:
    def test_h1_h2_creates_one_req_per_h2(self):
        data = _make_docx_bytes([
            ("Heading 1", "Mandatory Requirements"),
            ("Heading 2", "Product Sync"),
            (None, "Sync products daily from ERP."),
            ("Heading 2", "Asset Archive"),
            (None, "Archive obsolete assets."),
            ("Heading 1", "Optional"),
            ("Heading 2", "Reporting"),
            (None, "Generate weekly report."),
        ])
        reqs = _parse_docx_requirements(data)
        assert len(reqs) == 3

    def test_h2_heading_becomes_category(self):
        data = _make_docx_bytes([
            ("Heading 1", "Section"),
            ("Heading 2", "Product Sync"),
            (None, "Description here."),
        ])
        reqs = _parse_docx_requirements(data)
        assert reqs[0].category == "Product Sync"

    def test_body_text_becomes_description(self):
        data = _make_docx_bytes([
            ("Heading 2", "Feature"),
            (None, "First sentence."),
            (None, "Second sentence."),
        ])
        reqs = _parse_docx_requirements(data)
        assert "First sentence" in reqs[0].description
        assert "Second sentence" in reqs[0].description

    def test_mandatory_flag_from_h1_parent(self):
        data = _make_docx_bytes([
            ("Heading 1", "Mandatory Requirements"),
            ("Heading 2", "Sync"),
            (None, "Sync data."),
            ("Heading 1", "Non-Mandatory"),
            ("Heading 2", "Reporting"),
            (None, "Report data."),
        ])
        reqs = _parse_docx_requirements(data)
        sync = next(r for r in reqs if r.category == "Sync")
        report = next(r for r in reqs if r.category == "Reporting")
        assert sync.mandatory is True
        assert report.mandatory is False

    def test_h1_only_creates_one_req_per_h1(self):
        data = _make_docx_bytes([
            ("Heading 1", "Feature A"),
            (None, "Description A."),
            ("Heading 1", "Feature B"),
            (None, "Description B."),
        ])
        reqs = _parse_docx_requirements(data)
        assert len(reqs) == 2
        assert reqs[0].category == "Feature A"
        assert reqs[1].category == "Feature B"

    def test_no_headings_creates_one_req_per_paragraph(self):
        data = _make_docx_bytes([
            (None, "Sync products daily."),
            (None, "Archive old assets."),
        ])
        reqs = _parse_docx_requirements(data)
        assert len(reqs) == 2
        assert reqs[0].category == "General"

    def test_source_target_default_to_unknown(self):
        data = _make_docx_bytes([("Heading 1", "Feature"), (None, "Desc.")])
        reqs = _parse_docx_requirements(data)
        assert reqs[0].source_system == "Unknown"
        assert reqs[0].target_system == "Unknown"

    def test_auto_req_id_starts_with_r(self):
        data = _make_docx_bytes([("Heading 1", "Feature"), (None, "Desc.")])
        reqs = _parse_docx_requirements(data)
        assert reqs[0].req_id.startswith("R-")
